from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_cells_paragraph(cells, pt=8.0):
    def set_paragraph_font(paragraph):
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(pt)

    def set_paragraph_alignment(paragraph):
        paragraph.alignment = 1
        set_paragraph_font(paragraph)

    for cell in cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            set_paragraph_alignment(paragraph)


def create_table(page_number, table):
    set_table_border(table)

    title_row_cells = table.rows[0].cells
    title_row_cells[0].text = "当前页"
    title_row_cells[1].text = "今天"
    title_row_cells[2].text = "睡前"
    title_row_cells[3].text = "早上"
    title_row_cells[4].text = "第 2 天"
    title_row_cells[5].text = "第 3 天"
    title_row_cells[6].text = "第 5 天"
    title_row_cells[7].text = "第 8 天"
    title_row_cells[8].text = "第 15 天"
    title_row_cells[9].text = "第 30 天"
    set_cells_paragraph(title_row_cells, pt=6.5)

    index_row_cells = table.rows[1].cells
    index_row_cells[0].text = str(page_number)
    index_row_cells[1].text = str(page_number)
    index_row_cells[2].text = str(page_number)
    index_row_cells[3].text = str(page_number)
    index_row_cells[4].text = str(page_number + 1)
    index_row_cells[5].text = str(page_number + 2)
    index_row_cells[6].text = str(page_number + 4)
    index_row_cells[7].text = str(page_number + 7)
    index_row_cells[8].text = str(page_number + 14)
    index_row_cells[9].text = str(page_number + 29)
    set_cells_paragraph(index_row_cells, pt=10)


def set_table_vertical_position(table, position_cm):
    """设置表格在页面中的垂直位置"""
    tbl_pr = table._tbl.tblPr
    # 创建定位属性
    tblp_pr = OxmlElement("w:tblpPr")
    # 相对于页面定位
    tblp_pr.set(qn("w:vertAnchor"), "page")
    tblp_pr.set(qn("w:horzAnchor"), "margin")
    # 设置垂直位置（从页面顶部开始计算）
    tblp_pr.set(qn("w:tblpY"), str(int(position_cm * 360)))  # 1cm = 360 twips
    # 水平居中
    tblp_pr.set(qn("w:tblpXSpec"), "center")
    tbl_pr.append(tblp_pr)


for page in range(1, 7):
    table = doc.add_table(rows=2, cols=10)
    set_table_vertical_position(table, 36.9)

    create_table(page, table)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.add_page_break()


doc.save("memory_doc.docx")